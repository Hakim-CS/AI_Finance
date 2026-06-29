from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


OUT = Path(r"D:\Thesis\final\AI_Finance\codex\Aura_Finance_25_Page_Mechatronics_Body_Draft.docx")
CHART_DIR = Path(r"D:\Thesis\final\AI_Finance\AI\charts")


BODY = r"""
# 1 INTRODUCTION

## 1.1 Background and Motivation

Personal financial management is a continuous decision process rather than a static bookkeeping activity. Each month, a user receives income, observes expenditure across categories such as food, transport, utilities, health, shopping, entertainment, travel, and other costs, and attempts to allocate the remaining disposable income in a way that satisfies both spending needs and a saving target. Commercial finance applications record and visualize this activity, but most of them still rely on manually fixed budgets. Such budgets do not adapt when spending behavior changes, and they often fail to balance responsiveness with stability.

This thesis treats the budgeting problem as a software-based mechatronic control problem. In classical Mechatronics, sensors observe the state of a system, an estimator or model interprets the measurements, a controller computes an action, and a human-machine interface communicates the result to a user or operator. Aura Finance follows the same logic in a non-physical domain. Expense entries from manual input, voice recognition, and receipt extraction act as soft sensing channels. Monthly spending aggregation forms the observed state vector. A machine learning model estimates next-month spending, and an adaptive controller transforms that prediction into a feasible budget recommendation constrained by income and saving target.

The motivation for this framing is practical and academic. Practically, a recommendation that changes too sharply from month to month is difficult for a user to trust or follow. Academically, prediction accuracy alone is not enough to evaluate a control-oriented decision-support system. The output must also be stable, bounded, interpretable, and feasible under constraints. This thesis therefore evaluates budget allocation not only as a machine learning forecasting problem but also as a discrete-time control problem validated through software-in-the-loop simulation.

## 1.2 Research Question and Objectives

The central research question is: Can a machine-learning-assisted adaptive controller provide a better trade-off between budget tracking accuracy, overspending prevention, and budget stability than static budgeting or direct machine learning prediction alone?

The question deliberately asks for a trade-off rather than a single optimum. In a Mechatronics context, the best predictor is not necessarily the best controller. A controller must regulate the behavior of a system under uncertainty, disturbance, and constraints. In this thesis, the disturbance is monthly variability in spending behavior, the control action is the recommended category budget, and the feasibility constraint is the spendable income after the saving target has been reserved.

Table 1.1: Condensed research objectives.
| Ref. | Objective |
| --- | --- |
| O1 | Model personal budget allocation as a software mechatronic control loop with soft sensing, state estimation, prediction, control, and HMI components. |
| O2 | Generate synthetic persona-based spending data to support controlled software-in-the-loop evaluation without using private bank data. |
| O3 | Compare candidate regression models for monthly category-level spending prediction and select the model that best supports control use. |
| O4 | Design an adaptive budget controller that blends prediction with prior allocation and enforces income-saving feasibility. |
| O5 | Evaluate static, last-month, direct ML, and adaptive control strategies using tracking error, overspend rate, stability, and allocation accuracy. |

## 1.3 Scope of the Work

The thesis focuses on algorithmic budget allocation rather than complete commercial deployment. The web application is used as an implementation and HMI platform, while the main academic contribution lies in the control architecture and evaluation method. The dataset is synthetic and persona-based, so the results should be interpreted as controlled simulation evidence rather than proof of real-world behavioral generalization. The user is also not modeled as a behavioral feedback controller: the simulated future spending does not change in response to recommendations. This limitation is important, but it also keeps the experiment repeatable and suitable for software-in-the-loop comparison.

# 2 RELATED WORK AND RESEARCH GAP

## 2.1 Personal Finance Decision Support

Research on personal budgeting shows that budgets influence behavior, but budget compliance is weak when targets are optimistic or disconnected from observed behavior [1]. Digital finance tools can increase account monitoring, yet monitoring alone does not guarantee improved financial decisions [2]. This motivates the need for systems that do more than present historical spending. A useful budgeting assistant should transform measurements into actionable and stable recommendations.

Decision Support Systems provide a useful theoretical foundation. Sprague and Carlson describe decision support as the interaction of a data subsystem, a model subsystem, and a dialogue subsystem [3]. Aura Finance implements these three layers as a PostgreSQL-backed spending history, a Python prediction and control module, and a React dashboard for user interaction. Nunes and Jannach further emphasize that trust in algorithmic recommendations depends on explanations and consistency [4]. This directly supports the thesis focus on recommendation stability, not only predictive accuracy.

## 2.2 Machine Learning for Tabular Financial Prediction

The prediction component is based on structured tabular data: income, month number, previous category spending values, and total previous spending. For this type of dataset, tree-based ensemble methods are strong candidates because they capture nonlinear interactions without requiring very large training sets. Gradient Boosting follows a sequential residual-fitting principle in which each weak learner corrects errors from the previous ensemble [5]. Random Forests use bootstrap aggregation and feature subsampling to reduce variance [6]. Recent comparisons also show that tree-based models remain competitive on many tabular tasks where deep neural networks do not provide a practical advantage [14].

In this thesis, machine learning is not the final output. It is a predictive feedforward component inside a wider control architecture. The model estimates expected spending, but the controller decides how aggressively the budget should respond to that estimate. This separation is central to the Mechatronics framing because prediction and control have different performance goals.

## 2.3 Control, Soft Sensing, and Software-in-the-Loop Validation

Control theory describes how dynamic systems can be regulated under uncertainty, disturbance, and constraint [8]. Although the spending system is not a physical machine, its decision structure resembles many mechatronic systems: measurements arrive through sensors, the state is estimated, a controller computes an action, and the effect is observed in the next cycle. Exponential smoothing is used as a first-order discrete-time controller because it introduces inertia against sudden prediction changes while preserving adaptability.

Soft sensors are common in Mechatronics and process control when a desired variable cannot be measured directly or economically [12]. Aura Finance uses the same concept in a software environment. Voice input, receipt OCR, and manual expense forms are not physical sensors, but they transform user behavior into structured measurements. The monthly spending vector is therefore a soft-sensed state representation. The dashboard and alert system form the HMI, supporting situation awareness by showing the user state, constraints, and recommended action [13].

Software-in-the-loop validation is appropriate when the controller and software architecture must be tested before deployment to real users or physical systems [9]. It allows repeatable experiments, controlled disturbances, and direct comparison of strategies. Sargent argues that simulation models should be verified and validated against their intended use [10]. In this thesis, the intended use is not to predict every future expense perfectly, but to compare budget allocation strategies under the same synthetic spending conditions.

Table 2.1: Research gap motivating the condensed thesis.
| Area | Typical focus | Gap addressed in this thesis |
| --- | --- | --- |
| Personal finance apps | Tracking, categorization, visualization, manual budgets | Limited adaptive control of budget recommendations. |
| ML forecasting | One-step prediction accuracy | Recommendation stability and feasibility are usually secondary. |
| Decision support | Data, model, and interface integration | Control-system behavior is rarely evaluated. |
| Mechatronics/control | Physical systems and industrial automation | Software financial allocation can still be modeled as sensing, state estimation, control, and HMI. |
| SIL simulation | Controller validation before deployment | Useful for repeatable comparison of budgeting strategies. |

# 3 MECHATRONIC SYSTEM MODEL OF AURA FINANCE

## 3.1 System-Level Interpretation

The strongest Mechatronics interpretation of Aura Finance is not that money itself is a physical plant. Rather, the system follows the architecture of a mechatronic decision-support loop. A human user produces spending events. Software sensing channels convert these events into structured transaction records. Aggregation converts the transaction stream into a monthly state vector. A predictive model estimates future category demand. The controller computes a constrained budget allocation. The HMI presents the recommendation and budget status to the user.

This interpretation places the thesis inside the interdisciplinary territory of Mechatronics: measurement, information processing, control logic, software implementation, and human interaction. The control output is not a motor torque or valve position, but it is still an actionable command: a recommended spending limit for each category. The system must be responsive enough to adapt to changing spending, but stable enough that the user can plan around it.

Table 3.1: Mapping from Aura Finance components to Mechatronics concepts.
| Aura Finance component | Mechatronics concept | Role in the control loop |
| --- | --- | --- |
| Manual expense entry | Direct soft sensor | Captures user-provided amount, date, and category. |
| Voice expense input | Acoustic soft sensor | Converts spoken expense descriptions into structured records. |
| Receipt OCR input | Visual soft sensor | Extracts merchant and amount information from receipt images. |
| Monthly aggregation | State estimator | Forms the category-level spending state vector. |
| Gradient Boosting model | Predictive feedforward model | Estimates next-month category spending from the current state. |
| Saving target | Reference/constraint | Defines the spendable income ceiling. |
| Adaptive smoothing controller | Discrete-time controller | Balances responsiveness and stability under the income constraint. |
| Budget recommendation | Control action | Outputs feasible category-level spending limits. |
| Dashboard and alerts | Human-machine interface | Communicates state, recommendation, and risk to the user. |
| Next monthly spending | Feedback signal | Updates the measured state for the next cycle. |

## 3.2 State Vector and Signal Flow

The state representation is deliberately simple so that the control behavior can be interpreted. For each user and month, the feature vector contains income, month number, eight previous-month category spending values, and previous total spending. This gives an 11-dimensional state vector:

[[FORMULA: x_t = [income, month, prev_food, ..., prev_total]^T,  x_t in R^11     (3.1)]]

The target vector contains the observed monthly spending for eight categories. One model is trained per category, allowing the predictor to learn category-specific nonlinear patterns while keeping the architecture interpretable. The controller then normalizes its output so that the sum of all category budgets equals income minus the saving target.

The web application architecture supports this control loop. The frontend provides the HMI, the backend stores and manages measurements, and the Python module performs prediction, training, and software-in-the-loop evaluation. Figure 3.1 shows the database structure that supports the transaction and budget state. In the condensed thesis, implementation details are kept short because the Mechatronics contribution is the sensing-control-evaluation pipeline rather than the web interface itself.

[[FIGURE: ER diagram.png | Figure 3.1: Simplified data model supporting user state, transactions, categories, and budget recommendations. | 5.8]]

## 3.3 Controller Architecture

The controller is designed as a first-order adaptive allocation mechanism. Direct use of ML predictions would produce responsive but potentially volatile budgets. Static budgets would be stable but unable to adapt to changing behavior. The proposed controller occupies the middle ground: it smooths the ML estimate with the previous budget and then normalizes the result under the income-saving constraint.

The resulting control loop has three important properties. First, it is bounded because no category allocation can exceed the available spendable income after normalization. Second, it is adaptive because new predicted demand enters the controller at every month. Third, it is tunable through the smoothing parameter alpha. A low alpha makes the controller reactive; a high alpha makes it conservative. This tunable behavior is what allows the thesis to evaluate accuracy-stability trade-offs in control terms.

# 4 DATA GENERATION AND EXPERIMENTAL SETUP

## 4.1 Synthetic Digital-Twin Dataset

Real bank transaction data would raise privacy, consent, and reproducibility issues. The thesis therefore uses a synthetic persona-based dataset. In Mechatronics terms, the personas act as simplified digital twins of spending behavior. They do not represent exact real users, but they provide controlled state trajectories with known structure, seasonal variation, and category-dependent variability.

Five personas are simulated, ranging from a college student to an executive director. Each persona has a different income level, saving rate, and spending pattern. This gives the controller multiple operating regimes: low-income high-variability behavior, balanced professional spending, family-oriented spending, and high-income travel-heavy behavior. The aim is not to claim demographic realism, but to test whether the control strategy behaves consistently across different spending scales.

Table 4.1: Synthetic personas used as spending digital twins.
| Persona | Income | Saving rate | Dominant behavior |
| --- | --- | --- | --- |
| Emily, College Student | $2,800 | 5% | Food and entertainment dominant; high variability; minimal travel. |
| Marcus, Junior Developer | $5,500 | 15% | Balanced categories; moderate variability; occasional travel. |
| Sarah, Marketing Manager | $8,500 | 20% | Health and shopping emphasis; stable utilities; regular travel. |
| David, Senior Engineer | $13,000 | 25% | Family-oriented; consistent utilities and food; frequent travel. |
| Olivia, Executive Director | $22,000 | 30% | High travel and shopping; low relative food spend; low variability. |

## 4.2 Train-Test Protocol

The data are split directionally in time. Records from January 2023 to December 2024 are used for model training, and records from January 2025 to December 2025 are used for testing and software-in-the-loop control evaluation. This follows the operational direction of forecasting: the system must learn from past observations and make recommendations for future months. Random cross-validation would mix time periods and could produce overly optimistic estimates [7].

Table 4.2: Time-based experimental split.
| Split | Period | Rows | Purpose |
| --- | --- | --- | --- |
| Training | Jan 2023 - Dec 2024 | 115 | Fit all five regression candidates. |
| Test/SIL | Jan 2025 - Dec 2025 | 60 | Compare allocation strategies over 12 months and 5 personas. |

## 4.3 Evaluation Metrics

The evaluation uses metrics that reflect both prediction and control objectives. Tracking error measures how closely the recommended budget follows actual category spending. Overspend rate measures how often spending exceeds the recommended budget by more than 5 percent. Budget stability measures month-to-month change in recommendations. Allocation accuracy measures the directional agreement between recommended allocations and actual spending distribution.

[[FORMULA: TE = (1/N) sum |A_c,t - B_c,t|,  N = 480     (4.1)]]

[[FORMULA: OSR = count(A_c,t > 1.05 B_c,t) / 480     (4.2)]]

[[FORMULA: ST = (1/T) sum_t sum_c |B_c,t - B_c,t-1|     (4.3)]]

The four metrics intentionally represent competing objectives. A highly reactive strategy may achieve low tracking error but poor stability. A static strategy may have perfect stability but poor tracking. The adaptive controller is successful only if it improves the practical balance among these objectives.

# 5 PREDICTION AND ADAPTIVE CONTROL METHOD

## 5.1 Prediction Problem Formulation

The prediction task is monthly multi-output regression, implemented as separate category-specific regressors. For category c, the model estimates next-month spending from the state vector x_t:

[[FORMULA: y_hat_c,t = f_c(x_t)     (5.1)]]

Five candidate models are compared: Linear Regression, Decision Tree, Random Forest, Gradient Boosting, and Support Vector Regression. The comparison is included to justify the predictive component used by the controller, not to make the thesis primarily a machine learning benchmark. The selected model must provide sufficiently accurate predictions while remaining interpretable and stable enough for downstream control.

Table 5.1: Average model performance across all eight categories.
| Model | Avg MAE ($) | Avg RMSE ($) | Avg R2 | Train time |
| --- | --- | --- | --- | --- |
| Linear Regression | 114.21 | 161.48 | 0.9304 | 0.01 s |
| Decision Tree | 115.23 | 196.57 | 0.9021 | 0.02 s |
| Random Forest | 106.93 | 161.76 | 0.9332 | 0.55 s |
| Gradient Boosting | 90.98 | 150.96 | 0.9415 | 0.41 s |
| Support Vector Regression | 143.98 | 232.15 | 0.8561 | 0.02 s |

Gradient Boosting achieves the lowest average MAE, lowest average RMSE, and highest R2 among the tested models. The result supports its selection as the predictive feedforward model. Figure 5.1 keeps the model comparison in the condensed thesis, while detailed feature ablation and hyperparameter heatmaps can be moved to an appendix or supporting material.

[[FIGURE: model_comparison_mae.png | Figure 5.1: Average MAE comparison for the five candidate regression models. | 5.4]]

## 5.2 Budget Allocation Strategies

Four strategies are evaluated. Static Budget uses fixed category weights and therefore provides maximum stability but weak adaptability. Last-Month Baseline repeats the previous observed spending distribution. Direct ML Prediction normalizes predicted spending directly into a budget. Adaptive Controller blends the ML prediction with the previous final budget before normalization. This progression allows the experiment to isolate the value of prediction and then the value of control smoothing.

[[FORMULA: B_static,c = (I - S) w_c     (5.2)]]

[[FORMULA: B_LM,c,t = A_c,t-1 / sum_j A_j,t-1 * (I - S)     (5.3)]]

[[FORMULA: B_ML,c,t = y_hat_c,t / sum_j y_hat_j,t * (I - S)     (5.4)]]

[[FORMULA: B_adapt,c,t = alpha B_final,c,t-1 + (1 - alpha) B_ML,c,t     (5.5)]]

[[FORMULA: B_final,c,t = B_adapt,c,t / sum_j B_adapt,j,t * (I - S)     (5.6)]]

Equation (5.5) is the core control law. The smoothing parameter alpha determines how much inertia the controller has. Equation (5.6) enforces the budget feasibility constraint by ensuring that the sum of category budgets equals spendable income. This is the point at which the thesis becomes clearly control-oriented: the ML prediction is only an input signal, while the final decision is produced by a constrained controller.

## 5.3 Stability and Feasibility

The adaptive controller is stable in a bounded-input bounded-output sense. The direct ML budget B_ML,c,t is bounded by the spendable income ceiling I - S because it is normalized. The previous final budget B_final,c,t-1 was also bounded by the same normalization in the previous step. Since B_adapt,c,t is a convex combination of these bounded quantities for 0 < alpha < 1, it is also bounded:

[[FORMULA: 0 <= B_adapt,c,t <= I - S     (5.7)]]

The final normalization then guarantees feasibility across all categories. This prevents the controller from recommending a total budget that exceeds available income after saving. The proof is simple, but it is important because it connects the financial recommendation to control-system requirements: bounded output, constrained actuation, and predictable behavior under valid input conditions.

# 6 SOFTWARE-IN-THE-LOOP EVALUATION

## 6.1 Alpha Sensitivity

The smoothing parameter alpha is first evaluated to understand the controller trade-off. Low alpha values make the controller follow ML predictions closely, while high alpha values slow the response. Table 6.1 and Figure 6.1 show that alpha = 0.7 provides a strong reduction in instability while keeping tracking error within an acceptable range. The value is therefore selected for the final strategy comparison.

Table 6.1: Alpha sensitivity analysis.
| Alpha | Tracking error ($) | Stability ($) | OSR | Interpretation |
| --- | --- | --- | --- | --- |
| 0.1 | 182.62 | 822.41 | 24.2% | Near-ML accuracy but high volatility. |
| 0.3 | 185.94 | 638.67 | 24.4% | Modest smoothing; still volatile. |
| 0.5 | 192.99 | 464.63 | 23.8% | Balanced but less stable than selected setting. |
| 0.7 | 202.55 | 289.48 | 24.4% | Selected: strong stability reduction with acceptable accuracy loss. |
| 0.9 | 216.91 | 111.82 | 26.2% | Very stable but too slow to respond. |

[[FIGURE: control_alpha_sensitivity.png | Figure 6.1: Alpha sensitivity showing the tracking-error and stability trade-off. | 5.5]]

## 6.2 Strategy Comparison

The final software-in-the-loop experiment compares all four allocation strategies over the 2025 test period. The static strategy has zero month-to-month budget change, but its tracking error and overspend rate are high. The last-month baseline improves tracking but creates severe instability because it mirrors monthly spending swings. Direct ML prediction achieves the lowest tracking error, but its stability is poor for a user-facing recommendation system. The adaptive controller gives up some prediction-level accuracy in exchange for a large stability improvement.

Table 6.2: Main software-in-the-loop strategy comparison.
| Strategy | TE ($) | vs static | OSR | ST ($) | AA |
| --- | --- | --- | --- | --- | --- |
| Static Budget | 335.64 | - | 35.4% | 0.00 | 95.6% |
| Last-Month Baseline | 219.78 | -34.5% | 30.4% | 1,210.58 | 98.5% |
| Direct ML Prediction | 181.91 | -45.8% | 24.8% | 919.93 | 99.3% |
| Adaptive Controller | 202.55 | -39.7% | 24.4% | 289.48 | 98.7% |

[[FIGURE: control_comparison.png | Figure 6.2: Control-strategy comparison across tracking error, overspend rate, stability, and allocation accuracy. | 5.8]]

## 6.3 Accuracy-Stability Trade-Off

The key result is not that the adaptive controller has the lowest tracking error. It does not. Direct ML prediction is more accurate if the only objective is matching next-month spending. However, direct prediction changes sharply from month to month. The adaptive controller reduces budget instability from 919.93 dollars to 289.48 dollars relative to direct ML prediction, a reduction of 68.5 percent:

[[FORMULA: (919.93 - 289.48) / 919.93 = 68.5%     (6.1)]]

The cost of this stability improvement is an 11.3 percent increase in tracking error relative to direct ML prediction. For a Mechatronics-oriented decision-support system, this is an acceptable trade-off because the output is meant to guide human planning. A perfectly reactive budget can become a poor control output if it undermines trust, interpretability, or usability. Figure 6.3 visualizes the trade-off: the adaptive controller is positioned between the static and direct ML extremes.

[[FIGURE: control_tradeoff.png | Figure 6.3: Accuracy-stability trade-off among the evaluated allocation strategies. | 5.0]]

## 6.4 Interpretation of Results

The results answer the research question positively under the synthetic simulation conditions. A machine-learning-assisted adaptive controller provides a better practical control trade-off than static budgeting or direct ML prediction alone. It improves tracking relative to static budgeting, lowers overspending relative to static and last-month baselines, and substantially reduces volatility relative to direct ML. The result supports the thesis claim that adaptive budget allocation should be evaluated as a control problem, not only as a prediction problem.

Per-category error analysis shows that travel is the most difficult category because it is sporadic, seasonal, and high-amplitude. This is consistent with the control interpretation: categories with disturbance-like spikes are harder to regulate. Stable categories such as transport, health, and other are easier to predict and allocate. A future controller could therefore use category-specific alpha values, making stable categories more conservative and volatile categories more responsive.

# 7 DISCUSSION, LIMITATIONS, AND FUTURE WORK

## 7.1 Mechatronics Contribution

The main contribution of this thesis is the translation of adaptive budgeting into a mechatronic control-system structure. The work shows how soft sensing, state estimation, predictive modeling, constrained control, HMI design, and software-in-the-loop validation can be applied to a financial decision-support problem. This contribution is valuable because it expands Mechatronics beyond physical automation while preserving the discipline's core logic: observe, estimate, decide, act, and evaluate feedback behavior.

The thesis also demonstrates why control metrics are necessary. A conventional ML study would select the direct prediction model and stop at its MAE and R2. A Mechatronics-oriented study asks a different question: does the output behave well as a control action? The answer is more nuanced. Direct ML is accurate but volatile. Static budgeting is stable but inaccurate. The adaptive controller provides a regulated compromise, which is exactly the type of trade-off Mechatronics engineers routinely evaluate.

## 7.2 Limitations

The first limitation is the synthetic dataset. It provides repeatable experiments and privacy protection, but it cannot capture all real psychological, social, and economic factors that affect spending. The second limitation is the absence of behavioral feedback. In a real system, the user may change future spending in response to the recommendation, creating a true closed loop. The current simulation evaluates the controller against fixed spending trajectories rather than co-adaptive human behavior.

The third limitation is controller simplicity. Exponential smoothing is interpretable and stable, but it is only a first-order controller. More advanced approaches such as PID-style correction, Kalman filtering, or Model Predictive Control may improve the balance between tracking and stability. The fourth limitation is statistical scale: 60 test months across five personas are enough for controlled comparison, but not enough for strong statistical generalization to real users.

## 7.3 Future Work

Future work should first validate the method on anonymized real expense data. This would test whether the synthetic patterns used here transfer to actual user behavior. A second direction is category-adaptive smoothing, where alpha is tuned per category based on recent error variance. Stable categories could receive high alpha values, while volatile categories could respond faster. A third direction is behavioral closed-loop simulation, where a user model reacts to budget recommendations with partial compliance or delayed adaptation.

Further Mechatronics-oriented extensions include Model Predictive Control under savings and overspending constraints, Kalman-style state estimation for noisy or missing transactions, and HMI experiments measuring whether stable recommendations improve user trust and compliance. These extensions would strengthen the connection between software decision support and classical mechatronic design methodology.

# 8 CONCLUSION

This thesis presented Aura Finance as a software-based mechatronic control system for adaptive personal budget allocation. The system uses soft sensing channels to collect spending data, constructs a monthly spending state vector, applies machine learning to estimate future category demand, and uses an adaptive smoothing controller to produce feasible budget recommendations under an income-saving constraint.

The experimental results show that Gradient Boosting is the strongest predictive model among the tested candidates, achieving an average MAE of 90.98 dollars and R2 of 0.9415. However, the central finding is control-oriented rather than prediction-oriented. Direct ML prediction gives the lowest tracking error, but it produces volatile recommendations. The adaptive controller at alpha = 0.7 reduces instability by 68.5 percent relative to direct ML prediction while maintaining a large tracking improvement over static budgeting.

The research question is therefore answered affirmatively within the limits of the synthetic software-in-the-loop experiment. A machine-learning-assisted adaptive controller can provide a better practical trade-off between budget tracking accuracy, overspending prevention, and recommendation stability than static budgeting or direct ML prediction alone. The broader conclusion is that personal budgeting can be productively framed as a Mechatronics problem when the emphasis is placed on sensing, state estimation, constrained control, HMI, and validation under simulated operating conditions.

# REFERENCES AND LITERATURE

[1] M. F. Lukas and R. C. Howard, "The Influence of Budgets on Consumer Spending," Journal of Consumer Research, vol. 49, no. 5, pp. 697-720, 2023, doi: 10.1093/jcr/ucac024.

[2] S. Angel, "Smart tools? A randomized controlled trial on the impact of three different media tools on personal finance," Journal of Behavioral and Experimental Economics, vol. 74, pp. 104-111, 2018, doi: 10.1016/j.socec.2018.04.002.

[3] R. H. Sprague and E. D. Carlson, Building Effective Decision Support Systems. Englewood Cliffs, NJ: Prentice-Hall, 1982.

[4] I. Nunes and D. Jannach, "A systematic review and taxonomy of explanations in decision support and recommender systems," User Modeling and User-Adapted Interaction, vol. 27, pp. 393-444, 2017, doi: 10.1007/s11257-017-9195-0.

[5] J. H. Friedman, "Greedy function approximation: A gradient boosting machine," The Annals of Statistics, vol. 29, no. 5, pp. 1189-1232, 2001, doi: 10.1214/aos/1013203451.

[6] L. Breiman, "Random Forests," Machine Learning, vol. 45, pp. 5-32, 2001, doi: 10.1023/A:1010933404324.

[7] R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed. Melbourne, Australia: OTexts, 2021.

[8] K. J. Åström and R. M. Murray, Feedback Systems: An Introduction for Scientists and Engineers. Princeton, NJ: Princeton University Press, 2008.

[9] X. Chen, M. Salem, T. Das, and X. Chen, "Real Time Software-in-the-Loop Simulation for Control Performance Validation," Simulation, vol. 84, no. 8-9, pp. 399-407, 2008, doi: 10.1177/0037549708097420.

[10] R. G. Sargent, "Verification and validation of simulation models," Journal of Simulation, vol. 7, no. 1, pp. 12-24, 2013, doi: 10.1057/jos.2012.20.

[11] M. Tomizuka, "Mechatronics: from the 20th to 21st century," Control Engineering Practice, vol. 10, no. 8, pp. 877-886, 2002, doi: 10.1016/S0967-0661(02)00016-3.

[12] P. Kadlec, B. Gabrys, and S. Strandt, "Data-driven Soft Sensors in the process industry," Computers & Chemical Engineering, vol. 33, no. 4, pp. 795-814, 2009, doi: 10.1016/j.compchemeng.2008.12.012.

[13] M. R. Endsley, "Toward a theory of situation awareness in dynamic systems," Human Factors, vol. 37, no. 1, pp. 32-64, 1995, doi: 10.1518/001872095779049543.

[14] L. Grinsztajn, E. Oyallon, and G. Varoquaux, "Why tree-based models still outperform deep learning on tabular data," Advances in Neural Information Processing Systems, vol. 35, pp. 507-520, 2022.
"""


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(30)
    sec.bottom_margin = Mm(30)
    sec.left_margin = Mm(35)
    sec.right_margin = Mm(25)
    sec.header_distance = Mm(12.5)
    sec.footer_distance = Mm(12.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, bold, before, after in (
        ("Heading 1", 18, True, 0, 12),
        ("Heading 2", 14, False, 12, 6),
        ("Heading 3", 12, True, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    caption = doc.styles.add_style("Compact Caption", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    formula = doc.styles.add_style("Formula Text", 1)
    formula.font.name = "Consolas"
    formula.font.size = Pt(11)
    formula.paragraph_format.line_spacing = 1.15
    formula.paragraph_format.left_indent = Inches(0.65)
    formula.paragraph_format.space_before = Pt(4)
    formula.paragraph_format.space_after = Pt(4)
    return doc


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.add_run(text)


def add_formula(doc, text):
    p = doc.add_paragraph(style="Formula Text")
    p.add_run(text)


def add_caption(doc, text):
    p = doc.add_paragraph(style="Compact Caption")
    p.add_run(text)


def add_picture(doc, spec):
    parts = [p.strip() for p in spec.split("|")]
    filename, caption = parts[0], parts[1]
    width = float(parts[2]) if len(parts) > 2 else 5.5
    path = CHART_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        p.add_run().add_picture(str(path), width=Inches(width))
    else:
        p.add_run(f"[Missing figure file: {filename}]")
    add_caption(doc, caption)


def add_table(doc, caption, lines):
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    rows = [row for row in rows if not all(re.fullmatch(r"-+", cell.strip()) for cell in row)]
    if not rows:
        return
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for j, text in enumerate(rows[0]):
        cell = table.rows[0].cells[j]
        cell.text = text
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)
    for row in rows[1:]:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = text
            set_cell_margins(cells[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[j].paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("A Mechatronic Control-System Approach to Adaptive Budget Allocation Using Machine Learning")
    run.bold = True
    run.font.size = Pt(20)
    title.paragraph_format.space_after = Pt(12)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Condensed thesis body draft: Introduction to Conclusion")
    run.italic = True
    run.font.size = Pt(12)
    subtitle.paragraph_format.space_after = Pt(18)

    add_para(
        doc,
        "This draft condenses the original thesis body into a more Mechatronics-oriented structure of approximately 25 pages after final university formatting. It intentionally excludes the Slovenian-language requirements, UDC classification, language-editor information, and final template adjustments, which will be handled separately.",
    )

    lines = BODY.strip().splitlines()
    i = 0
    pending_caption = None
    first_h1 = True
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            i += 1
            continue
        if line.startswith("[[FORMULA:") and line.endswith("]]"):
            add_formula(doc, line[len("[[FORMULA:") : -2].strip())
            i += 1
            continue
        if line.startswith("[[FIGURE:") and line.endswith("]]"):
            add_picture(doc, line[len("[[FIGURE:") : -2].strip())
            i += 1
            continue
        if re.match(r"^Table \d+\.\d+:", line):
            pending_caption = line
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, pending_caption or "Table", table_lines)
            pending_caption = None
            continue
        if line.startswith("[") and re.match(r"^\[\d+\]", line):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Inches(-0.28)
            p.paragraph_format.left_indent = Inches(0.28)
            run = p.add_run(line)
            run.font.size = Pt(10.5)
            i += 1
            continue
        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("[[")
                or nxt.startswith("|")
                or re.match(r"^Table \d+\.\d+:", nxt)
                or re.match(r"^\[\d+\]", nxt)
            ):
                break
            para.append(nxt)
            i += 1
        add_para(doc, " ".join(para))

    for section in doc.sections:
        add_page_number(section)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
